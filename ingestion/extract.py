from __future__ import annotations

import logging
import os
import tempfile
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from ingestion.convert import convert_to_docx
from utils.validation import assert_allowed_filename

logger = logging.getLogger(__name__)

DOCX_SUFFIX = ".docx"
# Non-DOCX uploads are normalized to DOCX before extraction (PDF only from UI today).
CONVERT_TO_DOCX_SUFFIXES = frozenset({".pdf", ".doc", ".pptx"})


def extract_text(filename: str, content: bytes) -> tuple[str, dict]:
    """
    Normalize input to DOCX, then extract text from the DOCX structure.

    - DOCX: processed directly (no PDF conversion).
    - PDF (and other office types): converted to DOCX first, then processed as DOCX.
    """
    suffix = assert_allowed_filename(filename)
    try:
        docx_content, normalize_meta = _to_docx_bytes(filename, content, suffix)
    except Exception as exc:
        if suffix == ".pdf":
            logger.warning(
                "DOCX conversion failed for %s, falling back to PDF text extraction: %s",
                filename,
                exc,
            )
            text, extract_meta = _extract_from_pdf(content)
            extract_meta["extractor"] = "pdf_fallback"
            extract_meta["docx_error"] = str(exc)[:500]
            return text, {
                "source_format": "pdf",
                "normalized_format": "pdf",
                "converted_from": None,
                **extract_meta,
            }
        raise

    try:
        text, extract_meta = _extract_from_docx(docx_content)
    except Exception as exc:
        if suffix == ".pdf":
            logger.warning(
                "DOCX pipeline failed for %s, falling back to PDF text extraction: %s",
                filename,
                exc,
            )
            text, extract_meta = _extract_from_pdf(content)
            extract_meta["extractor"] = "pdf_fallback"
            extract_meta["docx_error"] = str(exc)[:500]
            return text, {**normalize_meta, **extract_meta}
        raise

    meta = {
        **normalize_meta,
        **extract_meta,
        "extractor": "docx",
    }
    return text, meta


def _to_docx_bytes(filename: str, content: bytes, suffix: str) -> tuple[bytes, dict]:
    if suffix == DOCX_SUFFIX:
        return content, {
            "source_format": "docx",
            "normalized_format": "docx",
            "converted_from": None,
        }

    if suffix not in CONVERT_TO_DOCX_SUFFIXES:
        raise ValueError(f"Cannot normalize {suffix} to DOCX.")

    with tempfile.TemporaryDirectory() as temp_dir:
        input_path = os.path.join(temp_dir, Path(filename).name)
        with open(input_path, "wb") as f:
            f.write(content)

        docx_path = convert_to_docx(input_path, temp_dir)
        with open(docx_path, "rb") as f:
            docx_bytes = f.read()

    source = suffix.lstrip(".")
    return docx_bytes, {
        "source_format": source,
        "normalized_format": "docx",
        "converted_from": source,
    }


def _extract_from_docx(content: bytes) -> tuple[str, dict]:
    doc = Document(BytesIO(content))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n\n".join(parts)
    if not text.strip():
        raise ValueError("DOCX contains no extractable text.")

    return text, {
        "type": "docx",
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "chars": len(text),
    }


def _extract_from_pdf(content: bytes) -> tuple[str, dict]:
    reader = PdfReader(BytesIO(content))
    parts = []
    for page in reader.pages:
        page_text = (page.extract_text() or "").strip()
        if page_text:
            parts.append(page_text)
    text = "\n\n".join(parts)
    if not text.strip():
        raise ValueError("PDF contains no extractable text.")
    return text, {
        "type": "pdf",
        "pages": len(reader.pages),
        "chars": len(text),
    }
